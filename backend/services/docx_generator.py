# backend/services/docx_generator.py

from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.shared import Pt

# base storage dir (like you do for pptx)
BASE_DIR = Path(__file__).resolve().parent.parent
DOC_STORAGE_DIR = BASE_DIR / "storage" / "docs"
DOC_STORAGE_DIR.mkdir(parents=True, exist_ok=True)


def _clean_section_content(doc_title: str, heading: str, raw: str) -> str:
    """
    Post-process Gemini text so that:

    - Leading lines like:
        * "Electric Vehicle Market in India 2025" (doc title)
        * "Page 1 – Section 1", "Page 2 - Section 3", etc.
        * "Section 2: Growth Drivers"
        * a repeat of the heading, or "Heading: ..."
      are removed (and we keep stripping such lines until the
      first "normal" paragraph).
    - Literal '\\n' are converted to real newlines.
    - Extra empty lines at the start/end are removed.
    """
    if not raw:
        return ""

    # Normalise newlines
    text = raw.replace("\r\n", "\n").replace("\r", "\n")

    # Convert escaped '\n' from the model into real paragraph breaks
    text = text.replace("\\n", "\n")

    lines = [l.strip() for l in text.split("\n")]

    doc_title_low = (doc_title or "").strip().lower()
    heading_low = (heading or "").strip().lower()

    # Strip meta lines at the TOP until we hit a normal paragraph
    while lines:
        first = lines[0].strip()
        if not first:
            # blank line at top – just drop it
            lines.pop(0)
            continue

        f_low = first.lower()

        is_doc_title = f_low == doc_title_low
        is_page_section = f_low.startswith("page ") and "section" in f_low
        is_section_prefix = f_low.startswith("section ")
        is_heading_exact = f_low == heading_low
        is_heading_with_colon = heading_low and f_low.startswith(heading_low + ":")

        if is_doc_title or is_page_section or is_section_prefix or is_heading_exact or is_heading_with_colon:
            lines.pop(0)
            continue

        # First "normal" line reached – stop stripping
        break

    # Collapse duplicate blank lines
    cleaned: List[str] = []
    for line in lines:
        if not line:
            if cleaned and cleaned[-1] == "":
                continue
            cleaned.append("")
        else:
            cleaned.append(line)

    # Trim blank lines at start/end
    while cleaned and cleaned[0] == "":
        cleaned.pop(0)
    while cleaned and cleaned[-1] == "":
        cleaned.pop()

    return "\n".join(cleaned).strip()


def build_docx_file(
    project_id: int,
    title: str,
    pages: Dict[int, List[Dict[str, str]]],  # {1: [...], 2: [...], ...}
) -> Path:
    """
    pages = {
      1: [
        {"heading": "Introduction", "content": "...."},
        {"heading": "Context", "content": "...."},
      ],
      2: [
        {"heading": "Market Analysis", "content": "...."},
      ],
      ...
    }

    - Each key is a page number (1-based).
    - Each value is a list of sections for that page.
    """

    doc = Document()

    # Title page (Word will handle its own pagination)
    doc.add_heading(title, level=0)

    first_page = True

    for page_num in sorted(pages.keys()):
        # For the first *content* page after the title we don't add a break.
        # For every later content page we insert an explicit page break.
        if not first_page:
            doc.add_page_break()
        first_page = False

        page_sections = pages[page_num]
        num_sections = len(page_sections)

        # Simple font size logic based on how many sections in the page
        if num_sections <= 1:
            para_size = Pt(12)
            heading_size = Pt(16)
        elif num_sections == 2:
            para_size = Pt(11)
            heading_size = Pt(14)
        else:
            # 3 sections -> slightly smaller text
            para_size = Pt(10)
            heading_size = Pt(13)

        for section in page_sections:
            heading = section.get("heading", "") or ""
            raw_content = section.get("content", "") or ""

            # Clean up Gemini text (remove doc title / Page X / Section X lines, handle \n)
            content = _clean_section_content(title, heading, raw_content)

            if heading:
                h = doc.add_heading(heading, level=1)
                for run in h.runs:
                    run.font.size = heading_size

            if content:
                for para_text in content.split("\n"):
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        for run in p.runs:
                            run.font.size = para_size

    file_name = f"project_{project_id}.docx"
    file_path = DOC_STORAGE_DIR / file_name
    doc.save(file_path)

    return file_path
